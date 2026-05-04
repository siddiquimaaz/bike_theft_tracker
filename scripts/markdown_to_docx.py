import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
MD_PATH = REPO_ROOT / "docs" / "TRD.md"
OUT_PATH = REPO_ROOT / "docs" / "TRD_print_ready.docx"


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_heading(document: Document, text: str, level: int) -> None:
    lvl = min(max(level, 1), 4)
    document.add_heading(text.strip(), level=lvl)


def add_image(document: Document, md_base: Path, src: str) -> None:
    clean_src = src.strip().split()[0]
    img_path = (md_base / clean_src).resolve()
    if img_path.exists():
        document.add_picture(str(img_path), width=Inches(6.2))
    else:
        document.add_paragraph(f"[Image not found: {clean_src}]")


def add_paragraph_with_bold(document: Document, text: str) -> None:
    # Very light markdown bold handling
    para = document.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return "|" in stripped and not stripped.startswith("[http")


def is_table_separator_row(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells if c)


def split_table_cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_markdown_table(document: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return

    rows = [split_table_cells(x) for x in table_lines if x.strip()]
    if len(rows) < 2:
        for line in table_lines:
            document.add_paragraph(line)
        return

    header = rows[0]
    body_rows = rows[1:]

    # Remove markdown separator row if present
    if body_rows and is_table_separator_row(table_lines[1]):
        body_rows = body_rows[1:]

    col_count = len(header)
    table = document.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header):
        run = hdr_cells[i].paragraphs[0].add_run(cell_text)
        run.bold = True

    for row in body_rows:
        row_cells = table.add_row().cells
        padded = row + [""] * max(0, col_count - len(row))
        for i in range(col_count):
            row_cells[i].text = padded[i]


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.splitlines()
    md_base = MD_PATH.parent

    document = Document()
    set_default_font(document)

    in_code_block = False
    code_buffer = []
    in_table_block = False
    table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_table_block:
                add_markdown_table(document, table_buffer)
                in_table_block = False
                table_buffer = []
            if in_code_block:
                if code_buffer:
                    p = document.add_paragraph("\n".join(code_buffer))
                    p.style = "No Spacing"
                in_code_block = False
                code_buffer = []
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if is_markdown_table_line(stripped):
            in_table_block = True
            table_buffer.append(stripped)
            continue
        elif in_table_block:
            add_markdown_table(document, table_buffer)
            in_table_block = False
            table_buffer = []

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith('<div style="page-break-after: always;"></div>'):
            document.add_page_break()
            continue

        if stripped == "---":
            document.add_paragraph("")
            continue

        if stripped.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m:
                level = len(m.group(1))
                add_heading(document, m.group(2), level)
                continue

        img_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if img_match:
            alt = img_match.group(1)
            src = img_match.group(2)
            if alt:
                cap = document.add_paragraph(alt)
                cap.runs[0].italic = True
            add_image(document, md_base, src)
            continue

        if stripped.startswith(("- ", "* ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^\d+\.\s+", "", stripped)
            document.add_paragraph(content, style="List Number")
            continue

        add_paragraph_with_bold(document, stripped)

    if in_table_block:
        add_markdown_table(document, table_buffer)

    document.save(OUT_PATH)
    print(f"Generated: {OUT_PATH}")


if __name__ == "__main__":
    main()
